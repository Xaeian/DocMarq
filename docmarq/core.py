# docmarq/core.py

"""Core DOCX class - main facade for document generation."""
import os
from typing import Callable
from docx import Document
from docx.shared import Mm, Pt, RGBColor, Emu
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .constants import Defaults, Align, A4
from .utils import (to_mm, parse_margin, align_to_docx, color_hex,
  rgb255, smaller_size, tight_line_height)
from .styles import Style, TableStyle
from .layout import PageGeometry
from .structure import Metadata
from .inline import RichSegment, _apply_run_format
from .tables import (
  apply_table_borders, apply_cell_shading, set_cell_align, repeat_header_row,
  set_cell_margins, set_cell_vertical_align,
)

#---------------------------------------------------------------------------------------- DOCX

class DOCX:
  """Main DOCX generator with fluent API.

  Example:
    >>> with DOCX("out.docx") as doc:
    ...   doc.heading("Tytuł", level=1)
    ...   doc.para("Pierwszy akapit.")
    ...   doc.text("Drugi z ").text("bold", bold=True).text(" fragmentem.")
  """
  def __init__(
    self,
    path: str,
    width: float = Defaults.PAGE_WIDTH,
    height: float = Defaults.PAGE_HEIGHT,
    margin: float|tuple = Defaults.MARGIN,
    unit: str = Defaults.UNIT,
    template: str|None = None,
    neutral_style: bool = True,
  ):
    self.path = path
    self.unit = unit
    top, right, bot, left = parse_margin(margin)
    self._page = PageGeometry(
      width=to_mm(width, unit), height=to_mm(height, unit),
      margin_top=to_mm(top, unit), margin_right=to_mm(right, unit),
      margin_bot=to_mm(bot, unit), margin_left=to_mm(left, unit),
    )
    self._doc = Document(template) if template else Document()
    self._apply_section_geometry()
    self._style = Style().with_defaults()
    self._metadata = Metadata()
    self._current_para = None # active `Paragraph` for run accumulation
    self._para_default_color = None # one-shot override; cleared on each `_flush_para`
    self._last_space_after = 0 # pt - previous block's space_after, used for
                                 # CSS-style margin collapsing on subsequent paragraphs
    self._extra_block_after = 0 # pt - one-shot bonus added to the NEXT paragraph's
                                 # `space_before` on top of collapse. Tables/lists set
                                 # this to force breathing room since their internal
                                 # paragraphs don't propagate `space_after` past the
                                 # block boundary.
    self._bookmark_id = 0
    if neutral_style:
      self._override_heading_styles()

  #---------------------------------------------------------------------------- Context manager

  def __enter__(self) -> "DOCX":
    return self

  def __exit__(self, exc_type, exc_val, exc_tb):
    if exc_type is None:
      self.save()

  #--------------------------------------------------------------------------------- Properties

  @property
  def page_width(self) -> float:
    """Current section's page width in mm."""
    return self._page.width

  @property
  def page_height(self) -> float:
    """Current section's page height in mm."""
    return self._page.height

  @property
  def content_width(self) -> float:
    """Drawable width in mm (page width minus left+right margins)."""
    return self._page.content_width

  @property
  def doc(self):
    """Direct `python-docx` `Document` access for advanced operations not
    exposed by the fluent API. Bypasses fluent state tracking."""
    return self._doc

  @property
  def output_path(self) -> str:
    """Destination file path. Symmetric with `PDF.output_path`."""
    return self.path

  #--------------------------------------------------------------------- Section / page setup

  def _override_heading_styles(self):
    """Replace Word's default blue Heading 1..6 with `pdfmarq`-style neutral
    palette - near-black text, GitHub-light sizes, tight spacing matching
    `pdfmarq`'s `head_gap_top`/`head_gap_bot`, thin bottom border for h1/h2.

    Word's stock heading colors come from the document theme (accent1).
    Setting `font.color.rgb` overrides the theme link so the color is fixed
    regardless of the active theme.
    """
    near_black = RGBColor(*rgb255(Defaults.HEAD_COLOR))
    rule_hex = color_hex(Defaults.RULE_COLOR)
    # Per-level (`space_before`, `space_after`) in pt - see `_HEADING_SPACING_PT`
    # on the class. Word ADDS adjacent vertical margins _(unlike CSS which
    # collapses them)_, so heading `space_before` is offset by previous block's
    # `space_after` in `heading()` itself - this style-level setting is just
    # the fallback default for paragraphs created outside the fluent API.
    spacing_pt = self._HEADING_SPACING_PT
    for i, size_pt in enumerate(Defaults.HEAD_SIZES, start=1):
      try:
        st = self._doc.styles[f"Heading {i}"]
      except KeyError:
        continue
      st.font.color.rgb = near_black
      st.font.size = Pt(size_pt)
      st.font.bold = True
      st.font.name = Defaults.FONT_FAMILY
      # `st.font.name` only sets `ascii` / `hAnsi`. The default Office theme
      # leaves `asciiTheme="majorHAnsi"` (= Calibri **Light**, not Calibri)
      # and some renderers prefer the theme reference over the direct font.
      # Strip the theme refs so we render with the stock body font everywhere.
      _force_font(st, Defaults.FONT_FAMILY)
      before_pt, after_pt = spacing_pt[i-1]
      st.paragraph_format.space_before = Pt(before_pt)
      st.paragraph_format.space_after = Pt(after_pt)
      # Display text needs tighter leading than body. `tight_line_height`
      # interpolates from 1.15 at body size down to 1.0 at 24pt+ so big
      # headings don't have wasteful vertical air between wrapped lines.
      st.paragraph_format.line_spacing = tight_line_height(size_pt)
      if i in Defaults.HEAD_UNDERLINE_LEVELS:
        _set_pbdr(st.element, rule_hex, sides=("bottom",), size_eighths=4)
    # Same treatment for `Normal` - guarantees body text uses Calibri across
    # all script ranges instead of any theme-resolved font.
    try:
      _force_font(self._doc.styles["Normal"], Defaults.FONT_FAMILY)
    except KeyError:
      pass

  def _apply_section_geometry(self):
    """Apply current `PageGeometry` to the first section."""
    sec = self._doc.sections[0]
    sec.page_width = Mm(self._page.width)
    sec.page_height = Mm(self._page.height)
    sec.top_margin = Mm(self._page.margin_top)
    sec.right_margin = Mm(self._page.margin_right)
    sec.bottom_margin = Mm(self._page.margin_bot)
    sec.left_margin = Mm(self._page.margin_left)
    sec.header_distance = Mm(self._page.effective_header_dist())
    sec.footer_distance = Mm(self._page.effective_footer_dist())
    if self._page.gutter:
      sec.gutter = Mm(self._page.gutter)

  def page(self, width:float|None=None, height:float|None=None) -> "DOCX":
    """Update page size for the current section."""
    if width: self._page.width = to_mm(width, self.unit)
    if height: self._page.height = to_mm(height, self.unit)
    self._apply_section_geometry()
    return self

  def margin(self, top:float|None=None, right:float|None=None,
      bot:float|None=None, left:float|None=None) -> "DOCX":
    """Update margins for the current section."""
    if top is not None: self._page.margin_top = to_mm(top, self.unit)
    if right is not None: self._page.margin_right = to_mm(right, self.unit)
    if bot is not None: self._page.margin_bot = to_mm(bot, self.unit)
    if left is not None: self._page.margin_left = to_mm(left, self.unit)
    self._apply_section_geometry()
    return self

  def landscape(self) -> "DOCX":
    """Swap page width/height for the current section."""
    self._page.width, self._page.height = self._page.height, self._page.width
    self._apply_section_geometry()
    return self

  #----------------------------------------------------------------------------------- Style

  def font(self, family:str|None=None, size:float|None=None,
      mode:str|None=None) -> "DOCX":
    """Set default font for subsequent runs.

    `mode` accepts `Regular` / `Bold` / `Italic` / `BoldItalic` for parity
    with `pdfmarq`; mapped to boolean `bold`/`italic` flags internally.
    """
    if family: self._style.font_family = family
    if size: self._style.font_size = size
    if mode:
      m = mode.lower()
      self._style.bold = "bold" in m
      self._style.italic = "italic" in m or "oblique" in m
    return self

  def color(self, color:tuple|str|None) -> "DOCX":
    """Set default text color for subsequent runs."""
    self._style.color = color
    return self

  def style(self, **overrides) -> "DOCX":
    """Update multiple style fields at once."""
    self._style = self._style.copy(**overrides)
    return self

  #--------------------------------------------------------------------------------- Paragraphs

  def para(self, text:str|None=None, style:str|None=None,
      align:str|None=None) -> "DOCX":
    """Start a new paragraph. Closes any previous active paragraph.

    Args:
      text: Optional initial text run.
      style: Built-in Word style name (e.g. `Quote`, `Intense Quote`).
      align: `L` / `R` / `C` / `J`.
    """
    self._flush_para()
    p = self._doc.add_paragraph(style=style)
    if align:
      p.alignment = align_to_docx(align)
    self._apply_para_spacing(p)
    self._current_para = p
    if text:
      self._add_run(p, RichSegment(text=text))
    return self

  def text(self, content:str, **seg_kwargs) -> "DOCX":
    """Add a styled run to the current paragraph. Auto-creates a paragraph
    if none active. Kwargs match `RichSegment` fields (`bold`, `italic`,
    `code`, `color`, `link_url`, etc.).
    """
    if self._current_para is None:
      self.para()
    seg = RichSegment(text=content, **seg_kwargs)
    self._add_run(self._current_para, seg)
    return self

  def link(self, content:str, url:str|None=None, target:str|None=None) -> "DOCX":
    """Add a hyperlink run. Either `url` for external link or `target` for
    internal bookmark name."""
    if self._current_para is None:
      self.para()
    self._add_hyperlink(self._current_para, content, url=url, target=target)
    return self

  def enter(self) -> "DOCX":
    """End the current paragraph - next `text()` starts a new one."""
    self._flush_para()
    return self

  def line_break(self) -> "DOCX":
    """Insert a soft line break inside the current paragraph (shift+enter)."""
    if self._current_para is None:
      self.para()
    run = self._current_para.add_run()
    run.add_break(WD_BREAK.LINE)
    return self

  def _flush_para(self):
    """Close the active paragraph - block elements must call this first."""
    self._current_para = None
    self._para_default_color = None

  def _apply_para_spacing(self, p):
    """Apply current `Style` line-height and before/after spacing to paragraph,
    with CSS-style margin collapsing against the previous block.

    Word natively SUMS adjacent paragraph margins _(unlike CSS which collapses
    them to `max(prev.after, this.before)`)_. We compensate by shrinking
    `this.space_before` by however much the previous block's `space_after`
    already contributes - visual gap = `max(prev_after, declared_before)`.
    On top of that, any `extra_block_after` left over from a heavy block
    (table, list) is added to the next paragraph's `space_before` since
    those blocks can't propagate `space_after` past their boundary.
    """
    pf = p.paragraph_format
    if self._style.line_height is not None:
      pf.line_spacing = self._style.line_height
    declared_before = self._style.space_before or 0
    declared_after = self._style.space_after or 0
    pf.space_before = Pt(self._collapsed_before(declared_before))
    pf.space_after = Pt(declared_after)
    self._last_space_after = declared_after

  def _collapsed_before(self, declared_before:float) -> float:
    """Compute effective `space_before`: collapse against previous block's
    `space_after`, then add any pending `extra_block_after`. Consuming this
    helper resets `_extra_block_after` so the bonus applies once only.
    """
    val = max(0, declared_before - self._last_space_after) + self._extra_block_after
    self._extra_block_after = 0
    return val

  def _track_block_spacing(self, after_pt:float, extra_for_next:float=0):
    """Record this block's `space_after` _(for collapsing against the next
    paragraph's `space_before`)_ plus an optional one-shot bonus added to
    the NEXT paragraph that goes through `_collapsed_before`.

    Tables and list runs use `extra_for_next` because their internal cell /
    item paragraphs have `space_after=0` (packed inside the block) and
    therefore don't produce any visible gap below the block on their own.
    """
    self._last_space_after = after_pt
    if extra_for_next:
      self._extra_block_after = extra_for_next

  def _add_run(self, p, seg:RichSegment):
    """Add a run with `seg` styling, falling back to current `_style`."""
    if seg.break_line:
      p.add_run().add_break(WD_BREAK.LINE)
    run = p.add_run(seg.text)
    # Effective family/size: segment wins, else current style
    family = seg.family or self._style.font_family
    size = seg.size or self._style.font_size
    # Color resolution priority: explicit segment color > per-paragraph default
    # (set by `blockquote()` etc.) > current `_style.color`.
    if seg.color is not None:
      eff_color = seg.color
    elif self._para_default_color is not None:
      eff_color = self._para_default_color
    else:
      eff_color = self._style.color
    # Style-level bold/italic OR-merged with segment-level
    merged = RichSegment(
      text=seg.text,
      family=family, size=size,
      color=eff_color,
      bold=seg.bold or bool(self._style.bold),
      italic=seg.italic or bool(self._style.italic),
      underline=seg.underline or bool(self._style.underline),
      strike=seg.strike or bool(self._style.strike),
      code=seg.code, highlight=seg.highlight or self._style.highlight,
      bg_color=seg.bg_color,
      superscript=seg.superscript,
      subscript=seg.subscript,
    )
    _apply_run_format(run, merged, family, size)

  #----------------------------------------------------------------------------------- Headings

  # Heading spacing per level (space_before, space_after) in pt. Mirrors
  # the values set on the style itself in `_override_heading_styles`. Used
  # both there AND in `heading()` per-paragraph to compute margin-collapsed
  # `space_before` against the previous block's `space_after`.
  _HEADING_SPACING_PT = [(10, 4), (8, 3), (6, 3), (4, 2), (3, 2), (3, 2)]

  def heading(self, text:str|None=None, level:int=1) -> "DOCX":
    """Add a heading (uses Word's built-in `Heading 1..9` styles for TOC support).
    Returns with `_current_para` set so `text()` can append rich runs to the
    heading - lets markdown render mixed inline formatting inside headings.

    Per-paragraph `space_before` is overridden to apply CSS-style margin
    collapsing against the previous block's `space_after` - prevents the
    "heading sits too far below body" effect Word otherwise produces.
    """
    self._flush_para()
    level = max(1, min(9, level))
    p = self._doc.add_paragraph(text or "", style=f"Heading {level}")
    spec_before, spec_after = self._HEADING_SPACING_PT[min(level - 1, 5)]
    eff_before = self._collapsed_before(spec_before)
    p.paragraph_format.space_before = Pt(eff_before)
    p.paragraph_format.space_after = Pt(spec_after)
    self._track_block_spacing(spec_after)
    self._current_para = p
    return self

  #-------------------------------------------------------------------------------------- Lists

  def bullet(self, text:str|None=None, level:int=0) -> "DOCX":
    """Bullet list item. Uses built-in `List Bullet` / `List Bullet 2..3` styles.

    For multi-level numbering with custom formats, define a numbering style
    via `doc.doc.styles` directly - this method covers the 99% case.
    """
    self._flush_para()
    style_name = "List Bullet" if level == 0 else f"List Bullet {level + 1}"
    try:
      p = self._doc.add_paragraph(text or "", style=style_name)
    except KeyError:
      p = self._doc.add_paragraph(text or "", style="List Bullet")
    self._tighten_list_spacing(p)
    # Items pack tightly internally (0/0) but the LAST item before a
    # non-list block must leave enough air for body to breathe. We can't
    # know which item is "last" until the next call lands, so every item
    # parks the same bonus - consumed only by `_collapsed_before`, which
    # the next list item bypasses (it goes through `_tighten_list_spacing`).
    self._track_block_spacing(0, extra_for_next=4)
    self._current_para = p
    return self

  def ordered(self, text:str|None=None, level:int=0) -> "DOCX":
    """Ordered list item. Uses built-in `List Number` / `List Number 2..3`."""
    self._flush_para()
    style_name = "List Number" if level == 0 else f"List Number {level + 1}"
    try:
      p = self._doc.add_paragraph(text or "", style=style_name)
    except KeyError:
      p = self._doc.add_paragraph(text or "", style="List Number")
    self._tighten_list_spacing(p)
    self._track_block_spacing(0, extra_for_next=4)
    self._current_para = p
    return self

  def _tighten_list_spacing(self, p):
    """Override Word's loose default list spacing - keep items packed."""
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.1

  #----------------------------------------------------------------------------- Block elements

  def code_block(self, content:str, language:str|None=None,
      bg_color:tuple|str=(0.96, 0.97, 0.98),
      border_color:tuple|str=(0.82, 0.84, 0.87),
      font_family:str="Consolas", font_size:float=9) -> "DOCX":
    """Insert a fenced code block - monospace font, light grey background,
    thin border. `language` is accepted but not used yet (no syntax highlight).
    """
    self._flush_para()
    p = self._doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(self._collapsed_before(2))
    pf.space_after = Pt(6)
    pf.line_spacing = 1.2
    _apply_paragraph_shading(p, color_hex(bg_color))
    _set_pbdr(p._element, color_hex(border_color), size_eighths=4)
    lines = content.rstrip("\n").split("\n")
    for i, line in enumerate(lines):
      if i > 0:
        p.add_run().add_break(WD_BREAK.LINE)
      run = p.add_run(line)
      run.font.name = font_family
      run.font.size = Pt(font_size)
    self._track_block_spacing(6)
    return self

  def blockquote(self, text:str|None=None,
      border_color:tuple|str=(0.82, 0.84, 0.87),
      text_color:tuple|str=(0.40, 0.44, 0.50),
      indent:float=4,
      space_before:float=3, space_after:float=3) -> "DOCX":
    """Insert a blockquote paragraph - thick left border, muted text, indent.

    For multi-paragraph blockquotes / callouts pass `space_before=0` on
    middle/last paragraphs and `space_after=0` on first/middle ones so the
    left bar reads as a single continuous block, not stacked stripes.

    Returns with `_current_para` set so `text()` can append rich runs.
    """
    self._flush_para()
    p = self._doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Mm(to_mm(indent, self.unit))
    pf.space_before = Pt(self._collapsed_before(space_before))
    pf.space_after = Pt(space_after)
    _set_pbdr(p._element, color_hex(border_color), sides=("left",), size_eighths=24, space=4)
    self._current_para = p
    self._para_default_color = text_color
    self._track_block_spacing(space_after)
    if text:
      self._add_run(p, RichSegment(text=text))
    return self

  def hr(self, color:tuple|str=(0.82, 0.84, 0.87), size_eighths:int=4) -> "DOCX":
    """Insert a horizontal rule - empty paragraph with bottom border."""
    self._flush_para()
    p = self._doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(self._collapsed_before(4))
    pf.space_after = Pt(4)
    _set_pbdr(p._element, color_hex(color), sides=("bottom",),
      size_eighths=size_eighths)
    self._track_block_spacing(4)
    return self

  #------------------------------------------------------------------------------------- Tables
  def table(
    self,
    body: list[list[str]],
    header: list[str]|None = None,
    aligns: list[str]|None = None,
    widths: list[float]|None = None, # mm per column
    style: TableStyle|None = None,
    word_style: str|None = None, # built-in Word table style name
  ) -> "DOCX":
    """Add a table.

    Args:
      body: List of rows (each row a list of cell strings).
      header: Optional header row.
      aligns: Per-column alignment - `L`/`R`/`C`/`J`.
      widths: Per-column widths in mm; `None` lets Word auto-size.
      style: `TableStyle` for borders/shading. Ignored when `word_style` set.
      word_style: Built-in Word style name (e.g. `Light Grid`).
    """
    self._flush_para()
    rows = ([header] if header else []) + (body or [])
    if not rows:
      return self
    ncols = max(len(r) for r in rows)
    tbl = self._doc.add_table(rows=len(rows), cols=ncols)
    if word_style:
      try: tbl.style = word_style
      except KeyError: pass
    style = style or TableStyle()
    # Effective cell font size: explicit `style.font_size` wins, else the
    # next-smaller value on the typographic ladder (11→10, 14→12, 16→14)
    # via `smaller_size`. Tables read better one ladder step below body.
    if style.font_size is None:
      body_pt = self._style.font_size or Defaults.FONT_SIZE
      eff_size_pt = smaller_size(body_pt)
    else:
      eff_size_pt = style.font_size
    # Column widths: explicit `widths=` wins; otherwise fill content area equally.
    # Without explicit widths Word/python-docx auto-sizes to content - making
    # the table appear left-shifted when the content is narrow. Setting explicit
    # widths (and disabling autofit) anchors the table to the full content area.
    if widths:
      col_w_mm = [to_mm(w, self.unit) for w in widths[:ncols]]
      while len(col_w_mm) < ncols:
        col_w_mm.append(self._page.content_width / ncols)
    elif style.fill_content_width:
      col_w_mm = [self._page.content_width / ncols] * ncols
    else:
      col_w_mm = None
    # Borders first - `_set_table_widths` below reorders `tblPr` children
    # into canonical schema order and will pick up these borders at the
    # right slot. Doing it the other way around leaves `tblBorders` at the
    # end of `tblPr`, which strict renderers (OnlyOffice) reject.
    apply_table_borders(tbl, style.border_color, style.border_size)
    if col_w_mm:
      tbl.autofit = False
      tbl.allow_autofit = False
      _set_table_widths(tbl, col_w_mm)
    # Table position on page
    if style.table_align:
      tbl.alignment = {"L": WD_TABLE_ALIGNMENT.LEFT, "C": WD_TABLE_ALIGNMENT.CENTER,
        "R": WD_TABLE_ALIGNMENT.RIGHT}.get(style.table_align, WD_TABLE_ALIGNMENT.LEFT)
    # Header row
    if header:
      hr = tbl.rows[0]
      for j, txt in enumerate(header):
        cell = hr.cells[j]
        cell.text = "" # clear placeholder
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(eff_size_pt)
        if style.header_bold:
          run.bold = True
        if style.header_color is not None:
          run.font.color.rgb = RGBColor(*rgb255(style.header_color))
        apply_cell_shading(cell, style.header_bg)
        set_cell_margins(cell, top=style.cell_pad_top, right=style.cell_pad_h,
          bot=style.cell_pad_bot, left=style.cell_pad_h)
        set_cell_vertical_align(cell, style.vertical_align)
        set_cell_align(cell, aligns[j] if aligns and j < len(aligns) else None)
        self._tighten_cell_paras(cell)
      if style.header_repeat:
        repeat_header_row(hr)
    # Body rows
    body_start = 1 if header else 0
    for i, row in enumerate(body or []):
      tr = tbl.rows[body_start + i]
      bg = style.row_bg_odd if i % 2 else style.row_bg_even
      for j, txt in enumerate(row):
        if j >= len(tr.cells):
          break
        cell = tr.cells[j]
        cell.text = "" # clear placeholder
        body_run = cell.paragraphs[0].add_run(str(txt))
        body_run.font.size = Pt(eff_size_pt)
        if bg is not None:
          apply_cell_shading(cell, bg)
        set_cell_margins(cell, top=style.cell_pad_top, right=style.cell_pad_h,
          bot=style.cell_pad_bot, left=style.cell_pad_h)
        set_cell_vertical_align(cell, style.vertical_align)
        set_cell_align(cell, aligns[j] if aligns and j < len(aligns) else None)
        self._tighten_cell_paras(cell)
    # Tables can't propagate `space_after` past their boundary _(cell paragraph
    # spacing stays inside the cell)_, so park a bonus for the next paragraph
    # to breathe. Without this, the body paragraph right below a table sticks
    # to its bottom border.
    self._track_block_spacing(0, extra_for_next=6)
    return self

  def _tighten_cell_paras(self, cell):
    """Strip Word's default body-paragraph spacing from cell content.
    Without this every cell carries phantom padding from `Normal` style's
    `space_after`, making rows look bottom-heavy regardless of `cell_pad_v`.
    """
    for p in cell.paragraphs:
      pf = p.paragraph_format
      pf.space_before = Pt(0)
      pf.space_after = Pt(0)
      pf.line_spacing = 1.15

  #------------------------------------------------------------------------------------ Images

  def image(self, path:str, width:float|None=None, height:float|None=None,
      align:str|None=None) -> "DOCX":
    """Insert an image. Width/height in current unit (mm by default).
    Provide only `width` for proportional scaling.
    """
    self._flush_para()
    p = self._doc.add_paragraph()
    if align:
      p.alignment = align_to_docx(align)
    run = p.add_run()
    kwargs = {}
    if width: kwargs["width"] = Mm(to_mm(width, self.unit))
    if height: kwargs["height"] = Mm(to_mm(height, self.unit))
    run.add_picture(path, **kwargs)
    # Treat image paragraph like a regular body block for collapse purposes.
    self._track_block_spacing(0)
    return self

  #------------------------------------------------------------------------------- Page breaks

  def page_break(self) -> "DOCX":
    """Insert a hard page break."""
    self._flush_para()
    p = self._doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return self

  #-------------------------------------------------------------------------------- Bookmarks

  def bookmark(self, name:str) -> "DOCX":
    """Mark a bookmark anchor at the current point (zero-width range).
    Use `link(..., target=name)` to jump here from another spot.
    """
    if self._current_para is None:
      self.para()
    self._bookmark_id += 1
    bid = str(self._bookmark_id)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)
    self._current_para._element.append(start)
    self._current_para._element.append(end)
    return self

  def _add_hyperlink(self, p, text:str, url:str|None=None, target:str|None=None):
    """Append a hyperlink run to paragraph. Either `url` or `target`."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    hl = OxmlElement("w:hyperlink")
    if url:
      r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
      hl.set(qn("r:id"), r_id)
    elif target:
      hl.set(qn("w:anchor"), target)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    # Hyperlink style: blue + underline
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rstyle)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    hl.append(run)
    p._element.append(hl)

  #-------------------------------------------------------------------------------- Headers/footers

  def header(self, text:str|None=None, align:str|None=None) -> "DOCX":
    """Set the header text for the current section.
    For richer headers use `doc.doc.sections[0].header` directly.
    """
    sec = self._doc.sections[-1]
    p = sec.header.paragraphs[0]
    p.text = text or ""
    if align: p.alignment = align_to_docx(align)
    return self

  def footer(self, text:str|None=None, align:str|None=None,
      page_number:bool=False) -> "DOCX":
    """Set the footer text for the current section.

    Args:
      text: Static footer content. Use `{page}` / `{pages}` for field codes.
      align: `L` / `R` / `C` / `J`.
      page_number: Append `Page X of Y` as Word fields (auto-updated).
    """
    sec = self._doc.sections[-1]
    p = sec.footer.paragraphs[0]
    p.text = ""
    if align: p.alignment = align_to_docx(align)
    if text:
      if "{page}" in text or "{pages}" in text:
        _inject_page_fields(p, text)
      else:
        p.add_run(text)
    if page_number:
      if text: p.add_run("  ")
      _inject_page_fields(p, "{page} / {pages}")
    return self

  #------------------------------------------------------------------------------------ Metadata

  def metadata(self, title:str|None=None, author:str|None=None,
      subject:str|None=None, keywords:str|None=None,
      comments:str|None=None, category:str|None=None) -> "DOCX":
    """Set core document properties."""
    if title: self._metadata.title = title
    if author: self._metadata.author = author
    if subject: self._metadata.subject = subject
    if keywords: self._metadata.keywords = keywords
    if comments: self._metadata.comments = comments
    if category: self._metadata.category = category
    return self

  #-------------------------------------------------------------------------------------- Save

  def save(self) -> "DOCX":
    """Write the document to disk."""
    self._flush_para()
    self._metadata.apply(self._doc)
    self._doc.save(self.path)
    return self

#--------------------------------------------------------------------------------- Page fields

def _inject_page_fields(p, template:str):
  """Insert Word `PAGE`/`NUMPAGES` fields where `{page}`/`{pages}` appear.

  Word's auto-page-numbering uses field codes (`<w:fldChar>` runs) rather
  than literal numbers. We split `template` on the placeholders and emit
  alternating literal-text runs and field-instruction runs.
  """
  import re
  parts = re.split(r"(\{page\}|\{pages\})", template)
  for part in parts:
    if part == "{page}":
      _append_field(p, "PAGE")
    elif part == "{pages}":
      _append_field(p, "NUMPAGES")
    elif part:
      r = p.add_run(part)

def _append_field(p, instr:str):
  """Append a `<w:fldSimple>` field run with given instruction."""
  fld = OxmlElement("w:fldSimple")
  fld.set(qn("w:instr"), instr)
  r = OxmlElement("w:r")
  t = OxmlElement("w:t")
  t.text = "0" # placeholder until Word recalculates
  r.append(t)
  fld.append(r)
  p._element.append(fld)

#---------------------------------------------------------------------- Paragraph shading/border

def _apply_paragraph_shading(p, hex_color:str):
  """Set background fill on a paragraph via raw `<w:shd>` in `pPr`."""
  pPr = p._element.get_or_add_pPr()
  shd = pPr.find(qn("w:shd"))
  if shd is None:
    shd = OxmlElement("w:shd")
    pPr.append(shd)
  shd.set(qn("w:val"), "clear")
  shd.set(qn("w:color"), "auto")
  shd.set(qn("w:fill"), hex_color)

_BORDER_SIDES_ALL = ("top", "left", "bottom", "right")

def _set_pbdr(element, hex_color:str, sides:tuple=_BORDER_SIDES_ALL,
    size_eighths:int=4, space:int=1):
  """Set paragraph borders on a paragraph or style element via `<w:pBdr>`.
  Works on both `<w:p>` and `<w:style>` since both expose `get_or_add_pPr()`.

  Args:
    element: `paragraph._element` or `style.element`.
    hex_color: Border color as `RRGGBB` hex string.
    sides: Subset of `("top", "left", "bottom", "right")`. Default: all four.
    size_eighths: Width in eighths-of-a-point (OOXML native unit).
    space: Gap from text in pt.
  """
  pPr = element.get_or_add_pPr()
  pBdr = pPr.find(qn("w:pBdr"))
  if pBdr is None:
    pBdr = OxmlElement("w:pBdr")
    pPr.append(pBdr)
  for side in sides:
    el = pBdr.find(qn(f"w:{side}"))
    if el is None:
      el = OxmlElement(f"w:{side}")
      pBdr.append(el)
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size_eighths))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), hex_color)

def _force_font(style, font_name:str):
  """Strip theme font references on a style and pin explicit `font_name`
  for every script range. Office's default theme has `majorHAnsi` resolving
  to "Calibri Light" (not Calibri), and some renderers prefer the theme
  reference over the direct `ascii` font - this guarantees Calibri wins.
  """
  from docx.oxml.ns import qn
  from docx.oxml import OxmlElement
  el = style.element
  rpr = el.find(qn("w:rPr"))
  if rpr is None:
    rpr = OxmlElement("w:rPr")
    el.append(rpr)
  rfonts = rpr.find(qn("w:rFonts"))
  if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
  # Drop every `*Theme` attribute - theme indirection is what causes the
  # Calibri-vs-Calibri-Light surprise.
  for theme_attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
    full = qn(f"w:{theme_attr}")
    if full in rfonts.attrib:
      del rfonts.attrib[full]
  # Pin direct fonts for every script range so Word/LibreOffice can't fall
  # back to a theme font we already stripped.
  for direct_attr in ("ascii", "hAnsi", "eastAsia", "cs"):
    rfonts.set(qn(f"w:{direct_attr}"), font_name)

#-------------------------------------------------------------------------------- Table width

# `dxa` (twentieths of a point) is the OOXML native length unit for tables.
_DXA_PER_MM = 1440 / 25.4

def _set_table_widths(tbl, col_widths_mm:list[float]):
  """Anchor a table's total width, per-column widths, left edge, and
  neutralize Word's implicit cell-margin default via raw XML.

  Children of `<w:tblPr>` are placed in canonical schema order
  (`tblW → jc → tblInd → tblBorders → tblLayout → tblCellMar → tblLook`).
  Strict OOXML renderers (OnlyOffice in particular) misinterpret tables
  whose `tblPr` children appear in non-schema order and silently fall back
  to defaults, which is why tables drifted left of the page margin there
  even though MS Word and LibreOffice rendered them correctly.

  `jc=left` is set explicitly: `tblInd` is documented as measured from
  the "leading edge", which is defined by `jc`. Without explicit `jc`,
  the leading edge is implementation-defined.

  `tblInd=0` anchors the table border at the page's left margin.

  `tblCellMar=0` neutralizes Word's table-level cell-margin default
  (typically 108 dxa ≈ 1.9 mm on left/right). Without this, Word extends
  the table border LEFT of `tblInd` to make room for the implicit default
  margin. Setting `tblCellMar=0` here lets per-cell `tcMar` fully control
  padding without disturbing the border position.
  """
  total_dxa = int(round(sum(col_widths_mm) * _DXA_PER_MM))
  tbl_el = tbl._element
  tbl_pr = tbl_el.find(qn("w:tblPr"))
  # Wipe pre-existing children we manage. Take references to elements we
  # want to keep BEFORE removing them so we can re-insert at the correct
  # schema slot.
  preserved_borders = tbl_pr.find(qn("w:tblBorders"))
  preserved_look = tbl_pr.find(qn("w:tblLook"))
  for tag in ("tblW", "jc", "tblInd", "tblBorders", "tblLayout",
      "tblCellMar", "tblLook"):
    for el in tbl_pr.findall(qn(f"w:{tag}")):
      tbl_pr.remove(el)
  # Build elements in canonical schema order, then append once at the end.
  # 1. tblW
  tbl_w = OxmlElement("w:tblW")
  tbl_w.set(qn("w:w"), str(total_dxa))
  tbl_w.set(qn("w:type"), "dxa")
  # 2. jc - explicit left justification (defines `tblInd` leading edge)
  jc = OxmlElement("w:jc")
  jc.set(qn("w:val"), "left")
  # 3. tblInd
  tbl_ind = OxmlElement("w:tblInd")
  tbl_ind.set(qn("w:w"), "0")
  tbl_ind.set(qn("w:type"), "dxa")
  # 4. tblBorders (carry over if any)
  # 5. tblLayout
  tbl_layout = OxmlElement("w:tblLayout")
  tbl_layout.set(qn("w:type"), "fixed")
  # 6. tblCellMar - zero on all sides via canonical `left`/`right`
  tbl_cell_mar = OxmlElement("w:tblCellMar")
  for side in ("top", "left", "bottom", "right"):
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:w"), "0")
    el.set(qn("w:type"), "dxa")
    tbl_cell_mar.append(el)
  # 7. tblLook (carry over if any)
  ordered = [tbl_w, jc, tbl_ind]
  if preserved_borders is not None:
    ordered.append(preserved_borders)
  ordered.append(tbl_layout)
  ordered.append(tbl_cell_mar)
  if preserved_look is not None:
    ordered.append(preserved_look)
  for el in ordered:
    tbl_pr.append(el)
  # Replace grid columns
  old_grid = tbl_el.find(qn("w:tblGrid"))
  if old_grid is not None:
    tbl_el.remove(old_grid)
  new_grid = OxmlElement("w:tblGrid")
  for w_mm in col_widths_mm:
    gc = OxmlElement("w:gridCol")
    gc.set(qn("w:w"), str(int(round(w_mm * _DXA_PER_MM))))
    new_grid.append(gc)
  # Insert tblGrid right after tblPr (OOXML schema requires this order).
  tbl_el.insert(list(tbl_el).index(tbl_pr) + 1, new_grid)
  # Also set each cell width so cell-level overrides are consistent.
  for row in tbl.rows:
    for i, w_mm in enumerate(col_widths_mm):
      if i < len(row.cells):
        row.cells[i].width = Mm(w_mm)
